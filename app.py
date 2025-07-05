import os
from flask import Flask, request, render_template, send_file
from deep_translator import GoogleTranslator
from werkzeug.utils import secure_filename
import docx
import fitz  # PyMuPDF
from docx import Document

# === Configuration ===
UPLOAD_FOLDER = 'uploads'
TRANSLATED_FOLDER = 'translated'
ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# === Utility Functions ===
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path, ext):
    if ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == 'docx':
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif ext == 'pdf':
        text = ''
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        return text
    return ''

def save_text(text, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(text)

# === Main Route ===
@app.route('/', methods=['GET', 'POST'])
def index():
    translated_text = None
    download_link = None

    if request.method == 'POST':
        lang = request.form.get('language')
        file_format = request.form.get('format', 'txt')  # Safe fallback
        file = request.files.get('file')

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            ext = filename.rsplit('.', 1)[1].lower()
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            original_text = extract_text(filepath, ext)
            translated_text = GoogleTranslator(source='auto', target=lang).translate(original_text)

            base_name = filename.rsplit('.', 1)[0]
            if file_format == 'docx':
                output_filename = f"translated_{base_name}.docx"
                output_path = os.path.join(TRANSLATED_FOLDER, output_filename)
                doc = Document()
                doc.add_paragraph(translated_text)
                doc.save(output_path)
            else:
                output_filename = f"translated_{base_name}.txt"
                output_path = os.path.join(TRANSLATED_FOLDER, output_filename)
                save_text(translated_text, output_path)

            download_link = f"/download/{output_filename}"

    return render_template('index.html', translated_text=translated_text, download_link=download_link)

# === File Download ===
@app.route('/download/<filename>')
def download_file(filename):
    path = os.path.join(TRANSLATED_FOLDER, filename)
    return send_file(path, as_attachment=True)

# === Start App ===
if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(TRANSLATED_FOLDER, exist_ok=True)
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
